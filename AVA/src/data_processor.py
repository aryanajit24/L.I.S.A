import os
import sys
import logging
import tempfile
from pathlib import Path
import json
from typing import Dict, List, Any, Union, Tuple, Optional
import time
import re
from datetime import datetime

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class DataProcessor:
    """
    Handles processing of different types of documents including:
    - Text documents (PDF, DOC, DOCX, TXT)
    - Images (JPG, PNG, BMP)
    - Videos (MP4, AVI, MOV)
    
    Provides methods to extract text, analyze content, and prepare data
    for the AVA model.
    """
    
    def __init__(self, config=None):
        """Initialize the data processor with available dependencies.
        
        Args:
            config: Optional configuration object for the data processor
        """
        self.config = config
        self.dependencies = {}
        self._check_dependencies()
        
        # Initialize cloud clients
        self.vision_client = None
        self.video_client = None
        self.speech_client = None
        self.gemini_client = None
        
        # Initialize cloud APIs if available
        self._init_cloud_apis()
        
        # Create necessary directories
        self._create_directories()
    
    def _check_dependencies(self) -> None:
        """Check which dependencies are available."""
        # Basic dependencies
        self.dependencies["PIL"] = self._check_import("PIL.Image", "Pillow")
        self.dependencies["numpy"] = self._check_import("numpy")
        self.dependencies["cv2"] = self._check_import("cv2", "opencv-python")
        
        # Document processing dependencies
        self.dependencies["PyMuPDF"] = self._check_import("fitz", "PyMuPDF")
        self.dependencies["pdf2image"] = self._check_import("pdf2image")
        
        # OCR dependencies
        self.dependencies["pytesseract"] = self._check_import("pytesseract")
        if self.dependencies["pytesseract"]:
            try:
                import pytesseract
                pytesseract.get_tesseract_version()
                self.dependencies["tesseract_installed"] = True
            except Exception:
                self.dependencies["tesseract_installed"] = False
                logger.warning("Tesseract OCR not properly installed or not in PATH")
                logger.warning("Install Tesseract OCR from: https://github.com/UB-Mannheim/tesseract/wiki")
        
        # Video processing dependencies
        self.dependencies["moviepy"] = self._check_import("moviepy.editor", "moviepy")
        if not self.dependencies["moviepy"]:
            logger.warning(
                "\n    MoviePy not properly configured. To use video processing features:"
                "\n    1. Install ImageMagick from https://imagemagick.org/script/download.php#windows"
                "\n    2. Add ImageMagick to your system PATH"
                "\n    3. Install moviepy: pip install moviepy"
                f"\n    Error: {sys.exc_info()[1] if sys.exc_info()[1] else 'No module named \'moviepy.editor\''}"
                "\n    "
            )
        
        # NLP dependencies
        self.dependencies["nltk"] = self._check_import("nltk")
        if self.dependencies["nltk"]:
            try:
                import nltk
                nltk.data.find('tokenizers/punkt')
                self.dependencies["nltk_data"] = True
            except (LookupError, ImportError):
                self.dependencies["nltk_data"] = False
                try:
                    nltk.download('punkt', quiet=True)
                    nltk.download('stopwords', quiet=True)
                    nltk.download('averaged_perceptron_tagger', quiet=True)
                    nltk.download('maxent_ne_chunker', quiet=True)
                    nltk.download('words', quiet=True)
                    self.dependencies["nltk_data"] = True
                except Exception:
                    pass
        
        # Cloud API dependencies
        self.dependencies["google_cloud_vision"] = self._check_import(
            "google.cloud.vision", "google-cloud-vision"
        )
        self.dependencies["google_cloud_videointelligence"] = self._check_import(
            "google.cloud.videointelligence", "google-cloud-videointelligence"
        )
        self.dependencies["google_cloud_speech"] = self._check_import(
            "google.cloud.speech", "google-cloud-speech"
        )
        self.dependencies["google_generativeai"] = self._check_import(
            "google.generativeai", "google-generativeai"
        )
        
        # Check for missing dependencies
        missing_packages = []
        core_dependencies = {
            "PyMuPDF": "PyMuPDF",
            "google_generativeai": "google-generativeai",
            "google_cloud_vision": "google-cloud-vision",
            "google_cloud_videointelligence": "google-cloud-videointelligence"
        }
        
        for dep, package in core_dependencies.items():
            if not self.dependencies.get(dep, False):
                missing_packages.append(package)
        
        if missing_packages:
            logger.error(f"Missing packages: {', '.join(missing_packages)}")
            logger.error(f"Install with: pip install {' '.join(missing_packages)}")
        
        # Check for OCR
        if not self.dependencies.get("tesseract_installed", False):
            logger.warning("Missing dependencies for full document processing:")
            logger.warning("Tesseract OCR")
            logger.warning("Install using:")
            logger.warning("pip install PyMuPDF pdf2image pytesseract")
            logger.warning("And install Tesseract OCR from: https://github.com/UB-Mannheim/tesseract/wiki")
    
    def _check_import(self, module_name: str, package_name: str = None) -> bool:
        """Try to import a module and return whether it's available."""
        try:
            __import__(module_name)
            return True
        except ImportError:
            return False
    
    def _init_cloud_apis(self) -> None:
        """Initialize cloud API clients if credentials are available."""
        if not os.environ.get("GOOGLE_APPLICATION_CREDENTIALS"):
            if os.path.exists("dogwood-boulder-458622-t1-839b69a572b8.json"):
                os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "dogwood-boulder-458622-t1-839b69a572b8.json"
        
        # Cloud Vision API
        if self.dependencies.get("google_cloud_vision", False):
            try:
                from google.cloud import vision
                self.vision_client = vision.ImageAnnotatorClient()
            except Exception as e:
                logger.warning(f"Could not initialize Vision API client: {e}")
        
        # Video Intelligence API
        if self.dependencies.get("google_cloud_videointelligence", False):
            try:
                from google.cloud import videointelligence
                self.video_client = videointelligence.VideoIntelligenceServiceClient()
            except Exception as e:
                logger.warning(f"Could not initialize Video Intelligence API client: {e}")
        
        # Speech API
        if self.dependencies.get("google_cloud_speech", False):
            try:
                from google.cloud import speech
                self.speech_client = speech.SpeechClient()
            except Exception as e:
                logger.warning(f"Could not initialize Speech API client: {e}")
        
        # Gemini API
        if self.dependencies.get("google_generativeai", False):
            try:
                import google.generativeai as genai
                # Check for a configuration method to verify the API
                if hasattr(genai, 'Client') or hasattr(genai, 'configure'):
                    self.gemini_client = genai
                else:
                    logger.warning("Could not initialize Gemini API: module 'google.generativeai' has no attribute 'Client'")
            except Exception as e:
                logger.warning(f"Could not initialize Gemini API: {e}")
    
    def _create_directories(self) -> None:
        """Create necessary directory structure for processing files."""
        base_dirs = [
            "data/processed/documents",
            "data/processed/images",
            "data/processed/videos",
            "data/raw/documents",
            "data/raw/images",
            "data/raw/videos",
            "temp"
        ]
        
        for dir_path in base_dirs:
            os.makedirs(dir_path, exist_ok=True)
    
    def process_document(self, file_path: str, doc_type: str = None) -> Dict[str, Any]:
        """
        Process a document file and extract relevant information.
        
        Args:
            file_path: Path to the document file
            doc_type: Type of document (e.g., "pdf", "image", "video")
                      If not provided, will be inferred from file extension
        
        Returns:
            Dictionary containing extracted data and metadata
        """
        if not os.path.exists(file_path):
            return {
                "error": f"File not found: {file_path}",
                "success": False
            }
        
        # Get file extension and infer document type if not provided
        _, ext = os.path.splitext(file_path)
        ext = ext.lower().lstrip(".")
        
        if not doc_type:
            doc_type = self._infer_doc_type(ext)
        
        logger.info(f"Processing {doc_type.upper()}: {file_path}")
        
        # Route to appropriate processing function based on document type
        if doc_type == "pdf":
            return self._process_pdf(file_path)
        elif doc_type == "image":
            return self._process_image(file_path)
        elif doc_type == "video":
            return self._process_video(file_path)
        elif doc_type in ["docx", "doc"]:
            return self._process_word_document(file_path)
        elif doc_type == "txt":
            return self._process_text_file(file_path)
        else:
            return {
                "error": f"Unsupported document type: {doc_type}",
                "success": False
            }
    
    def _infer_doc_type(self, extension: str) -> str:
        """Infer the document type based on file extension."""
        extension = extension.lower()
        
        if extension in ["pdf"]:
            return "pdf"
        elif extension in ["jpg", "jpeg", "png", "bmp", "gif", "tiff", "webp"]:
            return "image"
        elif extension in ["mp4", "avi", "mov", "wmv", "mkv", "flv", "webm"]:
            return "video"
        elif extension in ["doc", "docx"]:
            return "docx"
        elif extension in ["txt", "text"]:
            return "txt"
        else:
            return "unknown"
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Process a PDF document and extract text, images, and other content.
        Falls back to OCR if text extraction fails.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary containing extracted data and metadata
        """
        result = {
            "type": "document",
            "format": "pdf",
            "success": False,
            "metadata": {
                "filename": os.path.basename(file_path),
                "path": file_path,
                "size_bytes": os.path.getsize(file_path),
                "pages": 0,
                "created": datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
                "modified": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
            }
        }
        
        # Try using PyMuPDF (fitz) first
        if self.dependencies.get("PyMuPDF", False):
            try:
                import fitz  # PyMuPDF
                
                # Safely open the document
                doc = None
                try:
                    doc = fitz.open(file_path)
                    
                    # Get document info
                    result["metadata"]["pages"] = len(doc)
                    result["metadata"]["title"] = doc.metadata.get("title", "")
                    result["metadata"]["author"] = doc.metadata.get("author", "")
                    result["metadata"]["subject"] = doc.metadata.get("subject", "")
                    
                    # Extract text page by page
                    full_text = []
                    pages_text = {}
                    text_blocks = []
                    
                    for page_num, page in enumerate(doc):
                        try:
                            # Get page text
                            text = page.get_text()
                            full_text.append(text)
                            pages_text[str(page_num + 1)] = text
                            
                            # Get blocks for structure analysis
                            blocks = page.get_text("blocks")
                            for block in blocks:
                                if len(block) >= 5 and block[4]:  # Make sure it's a text block
                                    text_blocks.append({
                                        "text": block[4],
                                        "page": page_num + 1,
                                        "bbox": list(block[:4])
                                    })
                        except Exception as e:
                            logger.error(f"Error processing page {page_num + 1}: {e}")
                    
                    # Store extracted text
                    result["extracted_text"] = "\n".join(full_text)
                    result["pages_text"] = pages_text
                    result["text_blocks"] = text_blocks
                    result["success"] = True
                    
                except Exception as e:
                    logger.error(f"Error processing PDF: {e}")
                    if doc:
                        doc.close()
                    # Continue to fallback methods
                
                # Make sure to close the document
                if doc:
                    doc.close()
            
            except Exception as e:
                logger.error(f"Error with PyMuPDF: {e}")
        
        # If text extraction failed or no text was found, try pdf2image + OCR
        if not result.get("extracted_text") or not result.get("success", False):
            result = self._pdf_fallback_processing(file_path, result)
        
        # If we have extracted text, perform analysis
        if result.get("extracted_text") and result.get("success", False):
            try:
                analysis_result = self._analyze_text(result["extracted_text"])
                result["analysis"] = analysis_result
            except Exception as e:
                logger.error(f"Error in text analysis: {e}")
        
        # Add page count for tracking
        result["metadata"]["pages"] = result.get("metadata", {}).get("pages", 0)
        result["metadata"]["text_length"] = len(result.get("extracted_text", ""))
        
        return result
    
    def _pdf_fallback_processing(self, file_path: str, result: Dict[str, Any]) -> Dict[str, Any]:
        """Fallback method for PDF processing using OCR."""
        # Check if we have pdf2image and pytesseract
        if not self.dependencies.get("pdf2image", False) or not self.dependencies.get("tesseract_installed", False):
            logger.error("PDF fallback with pdf2image failed: Missing dependencies")
            return result
        
        try:
            from pdf2image import convert_from_path
            import pytesseract
            
            # Convert PDF to images
            images = convert_from_path(file_path)
            result["metadata"]["pages"] = len(images)
            
            # Process each image with OCR
            full_text = []
            pages_text = {}
            
            for i, image in enumerate(images):
                text = pytesseract.image_to_string(image)
                full_text.append(text)
                pages_text[str(i + 1)] = text
            
            # Store extracted text
            result["extracted_text"] = "\n".join(full_text)
            result["pages_text"] = pages_text
            result["success"] = True
            result["ocr_processed"] = True
            
            return result
            
        except Exception as e:
            logger.error(f"PDF fallback with pdf2image failed: {e}")
            return result
    
    def _process_image(self, file_path: str) -> Dict[str, Any]:
        """
        Process an image file using computer vision techniques.
        If Cloud Vision API credentials are available, use them for enhanced analysis.
        
        Args:
            file_path: Path to the image file
            
        Returns:
            Dictionary containing image analysis data
        """
        result = {
            "type": "image",
            "format": os.path.splitext(file_path)[1].lower().lstrip("."),
            "success": False,
            "metadata": {
                "filename": os.path.basename(file_path),
                "path": file_path,
                "size_bytes": os.path.getsize(file_path),
                "created": datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
                "modified": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
            }
        }
        
        # Process with PIL for basic metadata
        if self.dependencies.get("PIL", False):
            try:
                from PIL import Image
                with Image.open(file_path) as img:
                    result["metadata"]["dimensions"] = {
                        "width": img.width,
                        "height": img.height,
                    }
                    result["metadata"]["format"] = img.format
                    result["metadata"]["mode"] = img.mode
                    result["success"] = True
            except Exception as e:
                logger.error(f"Error processing image with PIL: {e}")
        
        # Process with OpenCV for more advanced analysis
        if self.dependencies.get("cv2", False):
            try:
                import cv2
                import numpy as np
                
                # Read image
                img = cv2.imread(file_path)
                if img is None:
                    raise ValueError("Could not read image with OpenCV")
                
                # Basic analysis
                result["opencv_analysis"] = {}
                
                # Convert to grayscale for analysis
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                
                # Calculate blurriness (lower is sharper)
                laplacian = cv2.Laplacian(gray, cv2.CV_64F)
                result["opencv_analysis"]["blurriness"] = 1.0 / (laplacian.var() + 1e-5)
                
                # Edge detection
                edges = cv2.Canny(gray, 100, 200)
                result["opencv_analysis"]["edge_percentage"] = np.count_nonzero(edges) / edges.size
                
                # Face detection
                face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
                if not face_cascade.empty():
                    faces = face_cascade.detectMultiScale(gray, 1.3, 5)
                    result["opencv_analysis"]["faces_detected"] = len(faces)
                    
                    if len(faces) > 0:
                        face_regions = []
                        for (x, y, w, h) in faces:
                            face_regions.append({
                                "x": int(x),
                                "y": int(y),
                                "width": int(w),
                                "height": int(h)
                            })
                        result["opencv_analysis"]["face_regions"] = face_regions
                
            except Exception as e:
                logger.error(f"Error processing image with OpenCV: {e}")
                result["opencv_analysis"] = {"error": str(e)}
        
        # Perform OCR with Tesseract if available
        if self.dependencies.get("tesseract_installed", False):
            try:
                import pytesseract
                from PIL import Image
                
                # Extract text using OCR
                img = Image.open(file_path)
                text = pytesseract.image_to_string(img)
                result["extracted_text"] = text
                
            except Exception as e:
                logger.error(f"Error during OCR: {e}")
        
        # Use Cloud Vision API for enhanced analysis if available
        if self.vision_client:
            try:
                from google.cloud import vision
                
                # Read image file
                with open(file_path, "rb") as image_file:
                    content = image_file.read()
                
                # Create image object
                image = vision.Image(content=content)
                
                # Perform detection
                response = self.vision_client.annotate_image({
                    'image': image,
                    'features': [
                        {'type_': vision.Feature.Type.TEXT_DETECTION},
                        {'type_': vision.Feature.Type.LABEL_DETECTION},
                        {'type_': vision.Feature.Type.FACE_DETECTION},
                        {'type_': vision.Feature.Type.LANDMARK_DETECTION},
                        {'type_': vision.Feature.Type.IMAGE_PROPERTIES}
                    ],
                })
                
                # Extract information
                vision_analysis = {}
                
                # Labels
                if response.label_annotations:
                    vision_analysis["labels"] = [
                        label.description for label in response.label_annotations
                    ]
                
                # Text
                if response.text_annotations:
                    vision_analysis["detected_text"] = response.text_annotations[0].description if response.text_annotations else ""
                
                # Faces
                if response.face_annotations:
                    vision_analysis["faces"] = len(response.face_annotations)
                    face_details = []
                    
                    for face in response.face_annotations:
                        face_detail = {
                            "joy": str(face.joy_likelihood),
                            "sorrow": str(face.sorrow_likelihood),
                            "anger": str(face.anger_likelihood),
                            "surprise": str(face.surprise_likelihood),
                            "under_exposed": str(face.under_exposed_likelihood),
                            "blurred": str(face.blurred_likelihood),
                            "headwear": str(face.headwear_likelihood)
                        }
                        face_details.append(face_detail)
                    
                    vision_analysis["face_details"] = face_details
                
                # Landmarks
                if response.landmark_annotations:
                    vision_analysis["landmarks"] = [
                        landmark.description for landmark in response.landmark_annotations
                    ]
                
                # Colors
                if response.image_properties_annotation:
                    colors = []
                    for color in response.image_properties_annotation.dominant_colors.colors:
                        colors.append({
                            "r": color.color.red,
                            "g": color.color.green,
                            "b": color.color.blue,
                            "score": color.score,
                            "pixel_fraction": color.pixel_fraction
                        })
                    vision_analysis["dominant_colors"] = colors
                
                # Objects
                if hasattr(response, "localized_object_annotations") and response.localized_object_annotations:
                    vision_analysis["objects"] = [
                        obj.name for obj in response.localized_object_annotations
                    ]
                
                # Add to result
                result["vision_analysis"] = vision_analysis
                
            except Exception as e:
                logger.error(f"Error with Cloud Vision API: {e}")
        
        return result
    
    def _process_video(self, file_path: str) -> Dict[str, Any]:
        """
        Process a video file to extract frames, audio, and other content.
        If Cloud Video API credentials are available, use for enhanced analysis.
        
        Args:
            file_path: Path to the video file
            
        Returns:
            Dictionary containing video analysis data
        """
        result = {
            "type": "video",
            "format": os.path.splitext(file_path)[1].lower().lstrip("."),
            "success": False,
            "metadata": {
                "filename": os.path.basename(file_path),
                "path": file_path,
                "size_bytes": os.path.getsize(file_path),
                "created": datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
                "modified": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
            }
        }
        
        # Get video metadata with OpenCV
        if self.dependencies.get("cv2", False):
            try:
                import cv2
                cap = cv2.VideoCapture(file_path)
                
                if not cap.isOpened():
                    logger.error(f"Could not open video file: {file_path}")
                else:
                    # Get video properties
                    fps = cap.get(cv2.CAP_PROP_FPS)
                    frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                    width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
                    height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
                    duration = frame_count / fps if fps > 0 else 0
                    
                    result["metadata"]["dimensions"] = {"width": width, "height": height}
                    result["metadata"]["fps"] = fps
                    result["metadata"]["frame_count"] = frame_count
                    result["metadata"]["duration"] = duration
                    
                    # Extract key frames
                    keyframes = []
                    max_keyframes = 5  # Limit to 5 keyframes
                    interval = max(1, frame_count // (max_keyframes + 1))
                    
                    for i in range(min(max_keyframes, frame_count)):
                        frame_pos = min((i + 1) * interval, frame_count - 1)
                        cap.set(cv2.CAP_PROP_POS_FRAMES, frame_pos)
                        ret, frame = cap.read()
                        if ret:
                            # Save frame to temp file
                            frame_filename = f"temp/keyframe_{i}.jpg"
                            cv2.imwrite(frame_filename, frame)
                            
                            # Process frame as image
                            frame_result = self._process_image(frame_filename)
                            
                            # Add to keyframes
                            keyframes.append({
                                "timestamp": frame_pos / fps,
                                "frame_number": frame_pos,
                                "vision_analysis": frame_result.get("vision_analysis", {})
                            })
                            
                            # Clean up
                            os.remove(frame_filename)
                    
                    result["keyframes"] = keyframes
                    result["success"] = True
                
                cap.release()
                
            except Exception as e:
                logger.error(f"Error processing video with OpenCV: {e}")
        
        # Process video with MoviePy if available
        if self.dependencies.get("moviepy", False):
            try:
                from moviepy.editor import VideoFileClip
                
                with VideoFileClip(file_path) as clip:
                    # Get basic properties
                    result["metadata"]["duration"] = clip.duration
                    result["metadata"]["fps"] = clip.fps
                    result["metadata"]["dimensions"] = {"width": clip.w, "height": clip.h}
                    
                    # Check audio
                    audio_analysis = {
                        "has_audio": clip.audio is not None
                    }
                    
                    if clip.audio:
                        audio_analysis["duration"] = clip.audio.duration
                        audio_analysis["fps"] = clip.audio.fps
                    
                    result["audio_analysis"] = audio_analysis
                    result["success"] = True
                
            except Exception as e:
                logger.error(f"Error processing video with MoviePy: {e}")
        
        # Use Video Intelligence API if available
        if self.video_client and os.environ.get("GOOGLE_APPLICATION_CREDENTIALS"):
            try:
                from google.cloud import videointelligence
                
                # Convert file to cloud storage or base64 content
                with open(file_path, "rb") as video_file:
                    content = video_file.read()
                
                # Note: For larger videos, we would upload to GCS first
                # but for simplicity, we'll use input content directly
                
                # Setup features to detect
                features = [
                    videointelligence.Feature.LABEL_DETECTION,
                    videointelligence.Feature.SHOT_CHANGE_DETECTION,
                    videointelligence.Feature.TEXT_DETECTION,
                ]
                
                # Add speech transcription if there's audio
                if result.get("audio_analysis", {}).get("has_audio", False):
                    features.append(videointelligence.Feature.SPEECH_TRANSCRIPTION)
                    speech_config = videointelligence.SpeechTranscriptionConfig(
                        language_code="en-US",
                        enable_automatic_punctuation=True,
                    )
                else:
                    speech_config = None
                
                # Analyze video content
                operation = self.video_client.annotate_video(
                    request={
                        "features": features,
                        "input_content": content,
                        "video_context": videointelligence.VideoContext(
                            speech_transcription_config=speech_config
                        ) if speech_config else None
                    }
                )
                
                logger.info("Video Intelligence API request sent, waiting for response...")
                
                # Wait for operation to complete (this can take some time)
                response = operation.result(timeout=180)  # 3-minute timeout
                
                # Process results
                annotation_results = response.annotation_results[0]
                
                cloud_analysis = {}
                
                # Labels
                if annotation_results.segment_label_annotations:
                    cloud_analysis["labels"] = []
                    for label in annotation_results.segment_label_annotations:
                        cloud_analysis["labels"].append({
                            "description": label.entity.description,
                            "confidence": label.segments[0].confidence
                        })
                
                # Shot changes
                if annotation_results.shot_annotations:
                    shots = []
                    for i, shot in enumerate(annotation_results.shot_annotations):
                        start_time = shot.start_time_offset.total_seconds()
                        end_time = shot.end_time_offset.total_seconds()
                        shots.append({
                            "scene_number": i + 1,
                            "start_time": start_time,
                            "end_time": end_time,
                            "duration": end_time - start_time
                        })
                    result["scene_analysis"] = shots
                
                # Text detection
                if annotation_results.text_annotations:
                    cloud_analysis["text_annotations"] = []
                    for text_annotation in annotation_results.text_annotations:
                        for segment in text_annotation.segments:
                            confidence = segment.confidence
                            start_time = segment.segment.start_time_offset.total_seconds()
                            end_time = segment.segment.end_time_offset.total_seconds()
                            cloud_analysis["text_annotations"].append({
                                "text": text_annotation.text,
                                "confidence": confidence,
                                "start_time": start_time,
                                "end_time": end_time
                            })
                
                # Speech transcription
                if hasattr(annotation_results, "speech_transcriptions") and annotation_results.speech_transcriptions:
                    cloud_analysis["speech_transcription"] = []
                    for transcription in annotation_results.speech_transcriptions:
                        for alternative in transcription.alternatives:
                            cloud_analysis["speech_transcription"].append({
                                "transcript": alternative.transcript,
                                "confidence": alternative.confidence
                            })
                
                # Add to result
                result["cloud_analysis"] = cloud_analysis
                
            except Exception as e:
                logger.error(f"Error with Video Intelligence API: {e}")
        
        return result
    
    def _process_word_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a Word document (.docx or .doc) to extract text and structure.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        result = {
            "type": "document",
            "format": os.path.splitext(file_path)[1].lower().lstrip("."),
            "success": False,
            "metadata": {
                "filename": os.path.basename(file_path),
                "path": file_path,
                "size_bytes": os.path.getsize(file_path),
                "created": datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
                "modified": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
            }
        }
        
        # Try to use python-docx if available
        try:
            import docx
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs]
            full_text = "\n".join(paragraphs)
            
            # Extract tables
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    "table_number": i + 1,
                    "rows": len(table.rows),
                    "columns": len(table.rows[0].cells) if table.rows else 0,
                    "data": table_data
                })
            
            # Extract headings
            headings = []
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    headings.append({
                        "level": int(para.style.name.replace('Heading', '')),
                        "text": para.text
                    })
            
            # Store results
            result["extracted_text"] = full_text
            result["tables"] = tables
            result["headings"] = headings
            result["success"] = True
            
            # Analyze text
            if full_text:
                result["analysis"] = self._analyze_text(full_text)
            
        except ImportError:
            logger.warning("python-docx not installed. Install with: pip install python-docx")
        except Exception as e:
            logger.error(f"Error processing Word document: {e}")
            
            # Fallback: try to convert to text using external tools or extract text from XML
            # This would be implemented if needed
        
        return result
    
    def _process_text_file(self, file_path: str) -> Dict[str, Any]:
        """
        Process a plain text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Dictionary containing text content and analysis
        """
        result = {
            "type": "document",
            "format": "txt",
            "success": False,
            "metadata": {
                "filename": os.path.basename(file_path),
                "path": file_path,
                "size_bytes": os.path.getsize(file_path),
                "created": datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
                "modified": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
            }
        }
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                # If all encodings failed, try binary mode
                with open(file_path, 'rb') as f:
                    text = f.read().decode('utf-8', errors='replace')
            
            result["extracted_text"] = text
            result["success"] = True
            
            # Analyze text
            if text:
                result["analysis"] = self._analyze_text(text)
                
        except Exception as e:
            logger.error(f"Error processing text file: {e}")
        
        return result
    
    def _analyze_text(self, text: str) -> Dict[str, Any]:
        """
        Analyze text content to extract statistics, key phrases, and entities.
        
        Args:
            text: Text content to analyze
            
        Returns:
            Dictionary containing text analysis
        """
        result = {}
        
        try:
            # Basic statistics
            lines = text.split('\n')
            words = text.split()
            
            # Count sentences with regex for better handling
            sentence_pattern = r'[.!?]+[\s\n]+|[.!?]+$'
            sentences = re.split(sentence_pattern, text)
            sentences = [s for s in sentences if s.strip()]
            
            # Estimate paragraphs (text blocks separated by newlines)
            paragraphs = [p for p in re.split(r'\n\s*\n', text) if p.strip()]
            
            # Basic statistics
            stats = {
                "character_count": len(text),
                "word_count": len(words),
                "sentence_count": len(sentences),
                "paragraph_count": len(paragraphs),
                "line_count": len(lines),
                "average_word_length": sum(len(word) for word in words) / max(len(words), 1),
                "average_sentence_length": len(words) / max(len(sentences), 1)
            }
            
            result["statistics"] = stats
            
            # Extract structure if NLTK is available
            if self.dependencies.get("nltk", True) and self.dependencies.get("nltk_data", True):
                import nltk
                from nltk.corpus import stopwords
                from nltk.tokenize import word_tokenize, sent_tokenize
                from nltk.tag import pos_tag
                from nltk.chunk import ne_chunk
                
                # Tokenize and remove stopwords
                stop_words = set(stopwords.words('english'))
                tokens = word_tokenize(text.lower())
                filtered_tokens = [word for word in tokens if word.isalnum() and word not in stop_words]
                
                # Get frequency distribution
                freq_dist = nltk.FreqDist(filtered_tokens)
                most_common = freq_dist.most_common(20)  # Top 20 words
                
                # Extract key phrases (n-grams)
                bigrams = list(nltk.bigrams(filtered_tokens))
                bigram_freq = nltk.FreqDist(bigrams)
                top_bigrams = bigram_freq.most_common(10)  # Top 10 bigrams
                
                # Format for output
                key_words = [word for word, count in most_common]
                key_phrases = [f"{w1} {w2}" for (w1, w2), count in top_bigrams]
                
                # Named entity recognition
                named_entities = []
                sentences = sent_tokenize(text)
                
                for sentence in sentences[:10]:  # Limit to first 10 sentences for performance
                    tokens = word_tokenize(sentence)
                    tagged = pos_tag(tokens)
                    entities = ne_chunk(tagged)
                    
                    # Extract named entities
                    for chunk in entities:
                        if hasattr(chunk, 'label'):
                            entity_text = ' '.join(c[0] for c in chunk)
                            entity_type = chunk.label()
                            named_entities.append({
                                "text": entity_text,
                                "type": entity_type
                            })
                
                # Add to result
                result["key_words"] = key_words
                result["key_phrases"] = key_phrases
                result["named_entities"] = named_entities
                
                # Identify potential document structure
                document_structure = {}
                
                # Identify potential headings
                potential_headings = []
                for i, line in enumerate(lines):
                    line = line.strip()
                    # Heading heuristics: short, ends with no period, followed by blank line
                    if (line and len(line) < 100 and 
                            not line.endswith('.') and 
                            (i+1 >= len(lines) or not lines[i+1].strip())):
                        potential_headings.append({
                            "line_number": i + 1,
                            "text": line
                        })
                
                # Identify potential lists
                list_markers = ['•', '-', '*', '1.', '2.', 'a.', 'b.', 'i.', 'ii.']
                lists = []
                current_list = None
                
                for i, line in enumerate(lines):
                    line = line.strip()
                    if any(line.startswith(marker) for marker in list_markers):
                        if current_list is None:
                            current_list = {"start_line": i + 1, "items": []}
                        current_list["items"].append(line)
                    elif current_list is not None and line == "":
                        lists.append(current_list)
                        current_list = None
                
                if current_list is not None:
                    lists.append(current_list)
                
                # Check for table of contents
                has_toc = False
                toc_patterns = ['table of contents', 'contents', 'index', 'chapters']
                lowercase_text = text.lower()
                
                for pattern in toc_patterns:
                    if pattern in lowercase_text:
                        # Look for patterns like "1. Introduction", "2. Methods", etc.
                        if re.search(r'\d+\.\s+\w+', text, re.MULTILINE):
                            has_toc = True
                            break
                
                # Add structure information
                document_structure["potential_headings"] = potential_headings
                document_structure["lists"] = lists
                document_structure["has_table_of_contents"] = has_toc
                
                result["document_structure"] = document_structure
                
                # Topic extraction (simple keyword-based approach)
                # In a production system, we might use more advanced topic modeling
                topic_keywords = {
                    "finance": ["money", "finance", "bank", "investment", "stock", "market", "financial", "economy"],
                    "technology": ["computer", "software", "hardware", "internet", "data", "algorithm", "technology", "ai"],
                    "health": ["health", "medical", "doctor", "patient", "disease", "hospital", "medicine", "treatment"],
                    "education": ["education", "school", "student", "learn", "teach", "university", "college", "academic"],
                    "environment": ["environment", "climate", "sustainability", "green", "eco", "pollution", "renewable"],
                    "business": ["business", "company", "corporate", "strategy", "management", "leadership", "organization"]
                }
                
                # Count topic keyword occurrences
                topic_counts = {topic: 0 for topic in topic_keywords}
                for word in filtered_tokens:
                    for topic, keywords in topic_keywords.items():
                        if word in keywords:
                            topic_counts[topic] += 1
                
                # Sort by count and get top topics
                sorted_topics = sorted(topic_counts.items(), key=lambda x: x[1], reverse=True)
                top_topics = [topic for topic, count in sorted_topics if count > 0][:3]  # Top 3 topics
                
                result["key_topics"] = top_topics
                
        except Exception as e:
            logger.error(f"Error in text analysis: {str(e)}")
            result["error"] = str(e)
        
        return result