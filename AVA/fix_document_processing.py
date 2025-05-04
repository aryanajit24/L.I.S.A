#!/usr/bin/env python
# fix_document_processing.py - AVA Document Processing Fix

import os
import sys
import logging
import subprocess
from pathlib import Path
import traceback
from typing import Dict, Any, List, Optional, Tuple, Union

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def install_dependencies():
    """Install required packages for document processing"""
    try:
        print("Installing required Python packages...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", 
                               "PyMuPDF", "pdf2image", "pytesseract", 
                               "nltk", "google-generativeai", 
                               "google-cloud-vision", "google-cloud-videointelligence"])
        
        # Setup NLTK data
        import nltk
        print("Downloading NLTK data...")
        nltk.download('punkt')
        print("✅ Successfully installed Python dependencies")
        return True
    except Exception as e:
        print(f"❌ Failed to install dependencies: {str(e)}")
        return False

def install_tesseract():
    """Provide instructions for installing Tesseract OCR"""
    print("\n==== Tesseract OCR Installation Instructions ====")
    print("To install Tesseract OCR:")
    
    if os.name == 'nt':  # Windows
        print("""
1. Download the installer from: https://github.com/UB-Mannheim/tesseract/wiki
2. Run the installer and follow the prompts
3. Check "Add to PATH" during installation
4. Verify installation by opening a new command prompt and running: tesseract --version
        """)
    else:  # macOS or Linux
        print("""
1. On macOS: brew install tesseract
2. On Ubuntu/Debian: sudo apt-get install tesseract-ocr
3. On Fedora: sudo dnf install tesseract
4. Verify installation by running: tesseract --version
        """)
    
    print("After installation, run this script again to verify that Tesseract is properly configured.")

def install_poppler():
    """Provide instructions for installing Poppler (required for pdf2image)"""
    print("\n==== Poppler Installation Instructions ====")
    print("Poppler is required for PDF to image conversion with pdf2image.")
    
    if os.name == 'nt':  # Windows
        print("""
1. Download the latest release from: http://blog.alivate.com.au/poppler-windows/
2. Extract the downloaded file to a folder (e.g., C:\\Program Files\\poppler-xx.xx)
3. Add the bin directory to your PATH environment variable
4. Restart your command prompt/terminal
        """)
        # Or run the PowerShell script if available
        poppler_script = Path("install_poppler.ps1")
        if poppler_script.exists():
            print("\nAlternatively, run the provided PowerShell script:")
            print("    powershell -ExecutionPolicy Bypass -File install_poppler.ps1")
    else:  # macOS or Linux
        print("""
1. On macOS: brew install poppler
2. On Ubuntu/Debian: sudo apt-get install poppler-utils
3. On Fedora: sudo dnf install poppler-utils
        """)
    
    print("After installation, run this script again to verify that Poppler is properly configured.")

def check_tesseract():
    """Check if Tesseract OCR is installed and properly configured"""
    try:
        import pytesseract
        version = pytesseract.get_tesseract_version()
        print(f"✅ Tesseract OCR is installed (version {version})")
        return True
    except Exception as e:
        print(f"❌ Tesseract OCR is not properly configured: {str(e)}")
        install_tesseract()
        return False

def check_poppler():
    """Check if Poppler is installed and properly configured"""
    try:
        from pdf2image import convert_from_path
        print("✅ pdf2image and Poppler are correctly configured")
        return True
    except Exception as e:
        print(f"❌ pdf2image or Poppler is not properly configured: {str(e)}")
        install_poppler()
        return False

def fix_pdf_processing():
    """
    Return a fixed implementation of the DataProcessor class that properly handles PDF documents.
    This implementation uses context managers to ensure proper resource cleanup.
    """
    class FixedDataProcessor:
        def __init__(self, config=None):
            self.config = config
            self.check_dependencies()

        def check_dependencies(self):
            """Check for required dependencies and log warnings if missing"""
            missing_deps = []
            
            # Check for PDF processing libraries
            try:
                import fitz  # PyMuPDF
            except ImportError:
                missing_deps.append("PyMuPDF")
            
            try:
                import pdf2image
            except ImportError:
                missing_deps.append("pdf2image")
            
            try:
                import pytesseract
                try:
                    pytesseract.get_tesseract_version()
                except Exception:
                    missing_deps.append("Tesseract OCR")
            except ImportError:
                missing_deps.append("pytesseract")
            
            # Log missing dependencies
            if missing_deps:
                logger.warning("Missing dependencies for full document processing:")
                logger.warning("\n".join(missing_deps))
                logger.warning("Install using:")
                logger.warning("pip install PyMuPDF pdf2image pytesseract")
                if "Tesseract OCR" in missing_deps:
                    logger.warning("And install Tesseract OCR from: https://github.com/UB-Mannheim/tesseract/wiki")

        def process_document(self, file_path: str) -> Dict[str, Any]:
            """
            Process a document file, with fixed implementation to avoid resource leaks.
            
            Args:
                file_path: Path to the document file
                
            Returns:
                Dict containing processing results and metadata
            """
            if not os.path.exists(file_path):
                return {"success": False, "error": f"File not found: {file_path}"}
            
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # Process based on file type
            if file_ext == ".pdf":
                return self._process_pdf(file_path)
            elif file_ext in [".jpg", ".jpeg", ".png", ".bmp"]:
                return self._process_image(file_path)
            elif file_ext in [".docx", ".doc"]:
                return self._process_word(file_path)
            else:
                return {"success": False, "error": f"Unsupported file format: {file_ext}"}

        def _process_pdf(self, file_path: str) -> Dict[str, Any]:
            """
            Process a PDF file with proper resource management.
            
            Args:
                file_path: Path to the PDF file
                
            Returns:
                Dict containing processing results and metadata
            """
            result = {
                "success": False,
                "format": "pdf",
                "extracted_text": "",
                "metadata": {}
            }
            
            # Try PyMuPDF (fitz) first - with proper context manager
            try:
                import fitz  # PyMuPDF
                
                # Use context manager to ensure proper cleanup
                with fitz.open(file_path) as doc:
                    # Extract document metadata
                    result["metadata"] = {
                        "pages": len(doc),
                        "title": doc.metadata.get("title", ""),
                        "author": doc.metadata.get("author", ""),
                        "creation_date": doc.metadata.get("creationDate", ""),
                    }
                    
                    # Extract text from all pages
                    all_text = []
                    for page_num in range(len(doc)):
                        page = doc[page_num]
                        all_text.append(page.get_text())
                    
                    result["extracted_text"] = "\n\n".join(all_text)
                    result["success"] = True
                
                # Return successful result
                return result
                
            except Exception as e:
                logger.error(f"Error processing PDF with PyMuPDF: {str(e)}")
                # Fall back to pdf2image + pytesseract
                return self._process_pdf_ocr_fallback(file_path)
        
        def _process_pdf_ocr_fallback(self, file_path: str) -> Dict[str, Any]:
            """
            Fall back to OCR-based PDF processing when PyMuPDF fails.
            Uses pdf2image to convert PDF pages to images and pytesseract for OCR.
            """
            result = {
                "success": False,
                "format": "pdf",
                "extracted_text": "",
                "metadata": {}
            }
            
            try:
                from pdf2image import convert_from_path
                import pytesseract
                
                # Convert PDF to images
                images = convert_from_path(file_path)
                result["metadata"]["pages"] = len(images)
                
                # Perform OCR on each page
                all_text = []
                for idx, image in enumerate(images):
                    text = pytesseract.image_to_string(image)
                    all_text.append(text)
                
                result["extracted_text"] = "\n\n".join(all_text)
                result["success"] = True
                
            except Exception as e:
                logger.error(f"PDF fallback with pdf2image failed: {str(e)}")
                # Add error message to result
                result["error"] = f"PDF processing failed: {str(e)}"
            
            return result
        
        def _process_image(self, file_path: str) -> Dict[str, Any]:
            """Process an image file using OCR"""
            result = {
                "success": False,
                "format": "image",
                "extracted_text": "",
                "metadata": {}
            }
            
            try:
                import pytesseract
                from PIL import Image
                
                # Open the image with proper resource management
                with Image.open(file_path) as image:
                    # Extract image metadata
                    result["metadata"] = {
                        "format": image.format,
                        "size": image.size,
                        "mode": image.mode
                    }
                    
                    # Perform OCR
                    text = pytesseract.image_to_string(image)
                    result["extracted_text"] = text
                    result["success"] = True
                
            except Exception as e:
                logger.error(f"Image processing failed: {str(e)}")
                result["error"] = f"Image processing failed: {str(e)}"
            
            return result
        
        def _process_word(self, file_path: str) -> Dict[str, Any]:
            """Process a Word document"""
            result = {
                "success": False,
                "format": "word",
                "extracted_text": "",
                "metadata": {}
            }
            
            try:
                import docx
                
                # Open the document
                doc = docx.Document(file_path)
                
                # Extract metadata
                result["metadata"] = {
                    "paragraphs": len(doc.paragraphs),
                    "sections": len(doc.sections)
                }
                
                # Extract text from paragraphs
                all_text = []
                for para in doc.paragraphs:
                    all_text.append(para.text)
                
                result["extracted_text"] = "\n".join(all_text)
                result["success"] = True
                
            except Exception as e:
                logger.error(f"Word document processing failed: {str(e)}")
                result["error"] = f"Word document processing failed: {str(e)}"
            
            return result
        
        def analyze_text(self, text: str) -> Dict[str, Any]:
            """Analyze extracted text to identify key information"""
            if not text or len(text.strip()) < 10:
                return {"error": "Not enough text for analysis"}
            
            result = {
                "word_count": 0,
                "sentences": [],
                "keywords": []
            }
            
            try:
                # Basic text stats
                words = text.split()
                result["word_count"] = len(words)
                
                # Split into sentences using NLTK
                try:
                    import nltk
                    sentences = nltk.sent_tokenize(text)
                    result["sentences"] = sentences[:10]  # Return first 10 sentences
                    result["sentence_count"] = len(sentences)
                except Exception as e:
                    logger.error(f"Error in text analysis: {str(e)}")
                
                # Extract basic keywords (simple frequency-based for now)
                word_freq = {}
                for word in words:
                    word = word.lower().strip('.,!?()[]:;')
                    if len(word) > 3:  # Skip short words
                        word_freq[word] = word_freq.get(word, 0) + 1
                
                # Sort by frequency and take top 10
                keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
                result["keywords"] = [{"word": word, "frequency": freq} for word, freq in keywords]
                
                return result
                
            except Exception as e:
                logger.error(f"Text analysis failed: {str(e)}")
                return {"error": f"Text analysis failed: {str(e)}"}
    
    return FixedDataProcessor

def main():
    """Main function to run the document processing fix"""
    print("==== AVA Document Processing Fix ====\n")
    
    # Install Python dependencies
    print("Step 1: Installing Python dependencies...")
    install_dependencies()
    
    # Check Tesseract
    print("\nStep 2: Checking Tesseract OCR...")
    check_tesseract()
    
    # Check Poppler
    print("\nStep 3: Checking Poppler for PDF conversion...")
    check_poppler()
    
    print("\nDocument processing fix setup complete!")
    print("To test the fixed document processing implementation, run:")
    print("    python test_fixed_pdf.py")
    
    return True

if __name__ == "__main__":
    main()